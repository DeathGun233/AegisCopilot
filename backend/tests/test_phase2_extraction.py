from __future__ import annotations

import sys
import zipfile
from io import BytesIO
from types import SimpleNamespace

from docx import Document as DocxDocument

from app.services.extraction import ExtractionService


def test_docx_extraction_preserves_tables_as_markdown_rows() -> None:
    document = DocxDocument()
    document.add_paragraph("欧洲 DDP 渠道清关资料表")
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "国家/地区"
    table.cell(0, 1).text = "渠道"
    table.cell(0, 2).text = "品类"
    table.cell(1, 0).text = "德国"
    table.cell(1, 1).text = "DDP空派"
    table.cell(1, 2).text = "带电"
    buffer = BytesIO()
    document.save(buffer)

    text = ExtractionService().extract("rules.docx", buffer.getvalue())

    assert "欧洲 DDP 渠道清关资料表" in text
    assert "表格 1" in text
    assert "| 国家/地区 | 渠道 | 品类 |" in text
    assert "| 德国 | DDP空派 | 带电 |" in text


def test_docx_extraction_reports_images_when_ocr_is_unavailable() -> None:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"></Types>',
        )
        archive.writestr("word/document.xml", "<w:document xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main'/>")
        archive.writestr("word/media/image1.png", b"not-a-real-image")

    text = ExtractionService().extract("image-only.docx", buffer.getvalue())

    assert "[图片 1：OCR 未启用或不可用]" in text


def test_pdf_extraction_preserves_page_numbers_and_table_like_rows(monkeypatch) -> None:
    class FakePage:
        def __init__(self, text: str) -> None:
            self._text = text
            self.images = []

        def extract_text(self) -> str:
            return self._text

    class FakeReader:
        def __init__(self, _stream) -> None:
            self.pages = [
                FakePage("国家/地区    渠道    品类\n德国    DDP空派    带电"),
                FakePage("版本\n生效日期：2026-05-01"),
            ]

    monkeypatch.setattr("app.services.extraction.PdfReader", FakeReader)

    text = ExtractionService().extract("rules.pdf", b"%PDF")

    assert "【第 1 页】" in text
    assert "【第 2 页】" in text
    assert "| 国家/地区 | 渠道 | 品类 |" in text
    assert "| 德国 | DDP空派 | 带电 |" in text
    assert "page_number: 1" not in text


def test_pdf_extraction_uses_optional_ocr_for_page_images(monkeypatch) -> None:
    class FakeImage:
        data = b"image-bytes"
        name = "image1.png"

    class FakePage:
        images = [FakeImage()]

        def extract_text(self) -> str:
            return ""

    class FakeReader:
        def __init__(self, _stream) -> None:
            self.pages = [FakePage()]

    fake_pillow = SimpleNamespace(open=lambda stream: stream)
    fake_tesseract = SimpleNamespace(image_to_string=lambda image, lang=None: "图片中的清关要求")
    monkeypatch.setattr("app.services.extraction.PdfReader", FakeReader)
    monkeypatch.setitem(sys.modules, "PIL", SimpleNamespace(Image=fake_pillow))
    monkeypatch.setitem(sys.modules, "PIL.Image", fake_pillow)
    monkeypatch.setitem(sys.modules, "pytesseract", fake_tesseract)

    text = ExtractionService(enable_ocr=True).extract("scan.pdf", b"%PDF")

    assert "图片 OCR 1：图片中的清关要求" in text
