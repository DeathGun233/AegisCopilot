from __future__ import annotations

import importlib
import zipfile
from io import BytesIO
from pathlib import Path

from .text import normalize_text

try:
    from pypdf import PdfReader
except ModuleNotFoundError:  # pragma: no cover - depends on local environment
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ModuleNotFoundError:  # pragma: no cover - depends on local environment
    DocxDocument = None


class ExtractionError(ValueError):
    pass


class ExtractionService:
    def __init__(self, *, enable_ocr: bool = False, ocr_languages: str = "chi_sim+eng") -> None:
        self.enable_ocr = enable_ocr
        self.ocr_languages = ocr_languages

    def extract(self, filename: str, content: bytes) -> str:
        suffix = Path(filename).suffix.lower()
        if suffix in {".txt", ".md", ".markdown"}:
            return self._decode_text(content)
        if suffix == ".pdf":
            return self._extract_pdf(content)
        if suffix == ".docx":
            return self._extract_docx(content)
        raise ExtractionError(f"暂不支持的文件类型：{suffix or '未知类型'}")

    @staticmethod
    def _decode_text(content: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
            try:
                return normalize_text(content.decode(encoding))
            except UnicodeDecodeError:
                continue
        raise ExtractionError("无法解码文本文件内容")

    def _extract_pdf(self, content: bytes) -> str:
        if PdfReader is None:
            raise ExtractionError("处理 PDF 需要先安装 pypdf")
        reader = PdfReader(BytesIO(content))
        page_blocks: list[str] = []
        for page_index, page in enumerate(reader.pages, start=1):
            lines = [f"【第 {page_index} 页】"]
            page_text = page.extract_text() or ""
            if page_text.strip():
                lines.append(_format_table_like_text(page_text))
            image_notes = self._extract_pdf_image_notes(page, page_index)
            lines.extend(image_notes)
            page_blocks.append("\n".join(lines))
        text = "\n\n".join(page_blocks)
        if not text.strip():
            raise ExtractionError("PDF 文本提取结果为空")
        return normalize_text(text)

    def _extract_docx(self, content: bytes) -> str:
        if DocxDocument is None:
            raise ExtractionError("处理 DOCX 需要先安装 python-docx")
        blocks: list[str] = []
        try:
            document = DocxDocument(BytesIO(content))
            blocks.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
            for table_index, table in enumerate(document.tables, start=1):
                rendered = _render_docx_table(table, table_index)
                if rendered:
                    blocks.append(rendered)
        except Exception:
            blocks = []
        blocks.extend(self._extract_docx_image_notes(content))
        text = "\n\n".join(blocks)
        if not text.strip():
            raise ExtractionError("DOCX 文本提取结果为空")
        return normalize_text(text)

    def _extract_pdf_image_notes(self, page, page_index: int) -> list[str]:
        notes: list[str] = []
        for image_index, image in enumerate(getattr(page, "images", []) or [], start=1):
            image_bytes = getattr(image, "data", b"") or b""
            ocr_text = self._ocr_image(image_bytes)
            if ocr_text:
                notes.append(f"图片 OCR {image_index}：{ocr_text}")
            else:
                notes.append(f"[第 {page_index} 页图片 {image_index}：OCR 未启用或不可用]")
        return notes

    def _extract_docx_image_notes(self, content: bytes) -> list[str]:
        notes: list[str] = []
        try:
            with zipfile.ZipFile(BytesIO(content)) as archive:
                image_names = sorted(name for name in archive.namelist() if name.startswith("word/media/"))
                for image_index, name in enumerate(image_names, start=1):
                    ocr_text = self._ocr_image(archive.read(name))
                    if ocr_text:
                        notes.append(f"图片 OCR {image_index}：{ocr_text}")
                    else:
                        notes.append(f"[图片 {image_index}：OCR 未启用或不可用]")
        except zipfile.BadZipFile:
            return []
        return notes

    def _ocr_image(self, image_bytes: bytes) -> str:
        if not self.enable_ocr or not image_bytes:
            return ""
        try:
            image_module = importlib.import_module("PIL.Image")
            pytesseract = importlib.import_module("pytesseract")
            image = image_module.open(BytesIO(image_bytes))
            return normalize_text(str(pytesseract.image_to_string(image, lang=self.ocr_languages)))
        except Exception:
            return ""


def _render_docx_table(table, table_index: int) -> str:
    rows = [[normalize_text(cell.text) for cell in row.cells] for row in table.rows]
    rows = [row for row in rows if any(cell for cell in row)]
    if not rows:
        return ""
    column_count = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (column_count - len(row)) for row in rows]
    header = normalized_rows[0]
    separator = ["---"] * column_count
    lines = [
        f"表格 {table_index}",
        _markdown_row(header),
        _markdown_row(separator),
    ]
    lines.extend(_markdown_row(row) for row in normalized_rows[1:])
    return "\n".join(lines)


def _format_table_like_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    table_blocks: list[str] = []
    index = 0
    while index < len(lines):
        cells = _split_table_like_line(lines[index])
        if len(cells) < 2:
            table_blocks.append(lines[index])
            index += 1
            continue
        table_rows = [cells]
        index += 1
        while index < len(lines):
            next_cells = _split_table_like_line(lines[index])
            if len(next_cells) != len(cells):
                break
            table_rows.append(next_cells)
            index += 1
        if len(table_rows) < 2:
            table_blocks.extend("    ".join(row) for row in table_rows)
            continue
        table_blocks.append(_markdown_row(table_rows[0]))
        table_blocks.append(_markdown_row(["---"] * len(table_rows[0])))
        table_blocks.extend(_markdown_row(row) for row in table_rows[1:])
    return "\n".join(table_blocks)


def _split_table_like_line(line: str) -> list[str]:
    return [cell.strip() for cell in line.split("    ") if cell.strip()]


def _markdown_row(cells: list[str]) -> str:
    return "| " + " | ".join(cell.replace("|", "\\|") for cell in cells) + " |"
